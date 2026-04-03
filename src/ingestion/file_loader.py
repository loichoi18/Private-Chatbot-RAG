"""
File Document Loader
Handles plain text files (.txt, .md, .csv, .docx, .html).
"""

import csv
import logging
from io import StringIO
from pathlib import Path

from langchain_core.documents import Document

from config.settings import settings

logger = logging.getLogger(__name__)


class FileLoader:
    """Load various file types into Documents."""

    LOADERS = {}  # extension -> method name

    def __init__(self):
        self.LOADERS = {
            ".txt": self._load_text,
            ".md": self._load_text,
            ".csv": self._load_csv,
            ".html": self._load_html,
            ".docx": self._load_docx,
        }

    def load(self, file_path: str) -> list[Document]:
        """Auto-detect file type and load."""
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = path.suffix.lower()
        loader_fn = self.LOADERS.get(ext)
        if loader_fn is None:
            raise ValueError(
                f"Unsupported file type: {ext}. "
                f"Supported: {list(self.LOADERS.keys())}"
            )

        logger.info(f"Loading {ext} file: {path.name}")
        return loader_fn(path)

    def _load_text(self, path: Path) -> list[Document]:
        text = path.read_text(encoding="utf-8")
        return [
            Document(
                page_content=text,
                metadata={
                    "source": str(path),
                    "filename": path.name,
                    "filetype": path.suffix,
                    "loader": "file",
                },
            )
        ]

    def _load_csv(self, path: Path) -> list[Document]:
        """Convert each row into a document."""
        text = path.read_text(encoding="utf-8")
        reader = csv.DictReader(StringIO(text))
        documents = []

        for i, row in enumerate(reader):
            content = "\n".join(f"{k}: {v}" for k, v in row.items() if v)
            if content.strip():
                documents.append(
                    Document(
                        page_content=content,
                        metadata={
                            "source": str(path),
                            "filename": path.name,
                            "row_index": i,
                            "loader": "csv",
                        },
                    )
                )
        return documents

    def _load_html(self, path: Path) -> list[Document]:
        from bs4 import BeautifulSoup

        html = path.read_text(encoding="utf-8")
        soup = BeautifulSoup(html, "html.parser")
        for tag in soup(["script", "style"]):
            tag.decompose()
        text = soup.get_text(separator="\n", strip=True)
        title = soup.find("title")

        return [
            Document(
                page_content=text,
                metadata={
                    "source": str(path),
                    "filename": path.name,
                    "title": title.string if title else "",
                    "loader": "html",
                },
            )
        ]

    def _load_docx(self, path: Path) -> list[Document]:
        from docx import Document as DocxDocument

        doc = DocxDocument(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        text = "\n\n".join(paragraphs)

        return [
            Document(
                page_content=text,
                metadata={
                    "source": str(path),
                    "filename": path.name,
                    "loader": "docx",
                    "paragraph_count": len(paragraphs),
                },
            )
        ]
